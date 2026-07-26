from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
DOCS.mkdir(exist_ok=True)

DOC_PATH = DOCS / "aiops_project_design_usage_20260520.docx"
IMG_PATH = DOCS / "aiops_architecture_20260520.png"


def font(size: int, bold: bool = False):
    candidates = [
        Path(r"C:\Windows\Fonts\msyhbd.ttc") if bold else Path(r"C:\Windows\Fonts\msyh.ttc"),
        Path(r"C:\Windows\Fonts\simhei.ttf"),
        Path(r"C:\Windows\Fonts\simsun.ttc"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def draw_architecture_diagram() -> None:
    width, height = 1800, 1120
    img = Image.new("RGB", (width, height), "#07111f")
    draw = ImageDraw.Draw(img)

    for x in range(0, width, 50):
        draw.line([(x, 0), (x, height)], fill="#0d2236", width=1)
    for y in range(0, height, 50):
        draw.line([(0, y), (width, y)], fill="#0d2236", width=1)
    draw.ellipse((1150, -180, 2050, 520), fill="#10224b")
    draw.ellipse((-300, 620, 620, 1420), fill="#072b3e")

    title_font = font(46, True)
    box_font = font(26, True)
    small_font = font(20)
    white = "#eaf6ff"
    cyan = "#28d7ff"
    amber = "#ffb347"
    purple = "#9f5cff"

    def rounded_box(xy, title, lines, border="#28d7ff", fill="#0b1830", title_color="#eaf6ff"):
        x1, y1, x2, y2 = xy
        draw.rounded_rectangle((x1 + 6, y1 + 8, x2 + 6, y2 + 8), radius=24, fill="#020712")
        draw.rounded_rectangle(xy, radius=24, fill=fill, outline=border, width=3)
        draw.text((x1 + 24, y1 + 18), title, font=box_font, fill=title_color)
        yy = y1 + 62
        for line in lines:
            draw.text((x1 + 26, yy), line, font=small_font, fill="#bcd3ef")
            yy += 31

    def arrow(p1, p2, color="#3cd8ff"):
        import math

        draw.line([p1, p2], fill=color, width=4)
        x1, y1 = p1
        x2, y2 = p2
        angle = math.atan2(y2 - y1, x2 - x1)
        length = 18
        a1 = angle + math.pi * 0.82
        a2 = angle - math.pi * 0.82
        points = [
            p2,
            (x2 + length * math.cos(a1), y2 + length * math.sin(a1)),
            (x2 + length * math.cos(a2), y2 + length * math.sin(a2)),
        ]
        draw.polygon(points, fill=color)

    draw.text((70, 45), "宁智网维 · 城域网 AIOps 总体架构", font=title_font, fill=white)
    draw.text(
        (72, 102),
        "江苏有线南京分公司｜数据接入、事件聚合、AI分析、规则降噪、Web演示控制台",
        font=small_font,
        fill="#99b6d6",
    )

    rounded_box((70, 190, 390, 390), "城域网设备", ["H3C / 核心路由交换", "Syslog UDP 10087", "SNMP Trap UDP 10086"], amber, "#10172a", amber)
    rounded_box((500, 150, 850, 430), "数据接入层", ["Logstash Syslog Pipeline", "Trap Pipeline", "MIB / OID 初步翻译", "原始数据完整留存"], cyan, "#081a2d", cyan)
    rounded_box((960, 110, 1310, 310), "Elasticsearch", ["syslog-raw / syslog-parsed", "trap-raw", "alarm-events", "高频时序检索与聚合"], cyan, "#081a2d", cyan)
    rounded_box((960, 360, 1310, 555), "MySQL 元数据", ["用户 / 角色 / 审计", "AI运行记录 / 发现项", "定时任务 / 规则 / 命中"], purple, "#11142d", purple)
    rounded_box((1420, 170, 1720, 450), "受控分析服务", ["Event Worker 聚合事件", "Current Window Summary", "Investigation Tools", "AI Scheduler / API"], cyan, "#081a2d", cyan)
    rounded_box((470, 615, 860, 870), "统一 LLM 调用", ["优先内网 deepseek-v4-pro", "失败自动 fallback 公网", "健康检查 / 耗时记录", "不输出 API Key"], amber, "#161527", amber)
    rounded_box((970, 650, 1325, 910), "AI分析与规则引擎", ["12小时/4小时窗口分析", "AI规则解析 + 后端兜底", "评分、降噪、安全例外", "结构化结论与报告"], purple, "#11142d", purple)
    rounded_box((1420, 660, 1720, 915), "Web 控制台", ["AI分析结果演示页", "实时 Events / Syslog / Trap", "AI分析规则管理", "历史、导出、用户管理"], cyan, "#081a2d", cyan)

    arrow((390, 290), (500, 290))
    arrow((850, 250), (960, 210))
    arrow((850, 335), (960, 445), purple)
    arrow((1310, 220), (1420, 265))
    arrow((1310, 455), (1420, 350), purple)
    arrow((1550, 450), (1140, 650), amber)
    arrow((860, 735), (970, 770), amber)
    arrow((1325, 780), (1420, 780))
    arrow((1540, 660), (1540, 450), cyan)

    draw.rounded_rectangle((70, 980, 1720, 1055), radius=18, fill="#08182b", outline="#1b7898", width=2)
    draw.text(
        (100, 1000),
        "设计边界：AI 不直接访问 ES/MySQL，不执行自由 SQL/DSL；后端工具受控取证，规则降噪必须保留重大故障安全例外。",
        font=small_font,
        fill="#d6ecff",
    )
    draw.text(
        (100, 1030),
        "部署边界：程序目录 /opt/jscn-aiops，数据目录 /data/jscn-aiops，配置与密钥通过 deploy/.env 注入。",
        font=small_font,
        fill="#d6ecff",
    )
    img.save(IMG_PATH)


def set_cell_bg(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = header
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_bg(hdr[i], "1F4E79")
        for paragraph in hdr[i].paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    return table


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_code(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="CodeBlock")
    paragraph.add_run(text)


def build_document() -> None:
    draw_architecture_diagram()

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    styles = doc.styles
    for name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3"]:
        styles[name].font.name = "Microsoft YaHei"
        styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)
    styles["Heading 1"].font.size = Pt(18)
    styles["Heading 2"].font.size = Pt(14)
    styles["Heading 3"].font.size = Pt(12)
    if "CodeBlock" not in styles:
        styles.add_style("CodeBlock", 1)
    styles["CodeBlock"].font.name = "Consolas"
    styles["CodeBlock"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["CodeBlock"].font.size = Pt(9)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("宁智网维 · 城域网 AI 运维中枢")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(21, 93, 143)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("AIOps 项目设计说明与使用说明")
    run.bold = True
    run.font.size = Pt(20)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("江苏有线南京分公司｜城域网智能运维分析演示系统")
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(80, 100, 120)
    for line in [
        "版本：V1.0",
        "日期：2026-05-20",
        "代码仓库：NJCATV/AIOps",
        "部署位置：/opt/jscn-aiops，数据目录：/data/jscn-aiops",
    ]:
        paragraph = doc.add_paragraph(line)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()
    doc.add_heading("文档说明", level=1)
    add_table(
        doc,
        ["项目", "说明"],
        [
            ["系统名称", "宁智网维 · 城域网 AI 运维中枢"],
            ["建设单位", "江苏有线南京分公司"],
            ["项目定位", "面向城域网的 AIOps 智能分析、故障研判、规则降噪和演示展示系统"],
            ["当前形态", "单机 Docker Compose 部署，具备后续迁移到多节点架构的配置边界"],
            ["主要用户", "网管值班人员、网络维护工程师、运维主管、系统管理员"],
            ["文档范围", "系统设计、数据流、AI分析流程、规则功能、部署配置、页面使用和运维检查"],
        ],
    )

    doc.add_heading("1. 项目概述", level=1)
    doc.add_paragraph(
        "宁智网维 · 城域网 AI 运维中枢是面向江苏有线南京分公司城域网运维场景建设的 AIOps 演示与实用系统。"
        "系统接入城域网设备上报的 Syslog 与 SNMP Trap，完成原始数据留存、结构化解析、事件聚合、AI辅助研判、自然语言规则降噪和可视化展示。"
    )
    doc.add_paragraph(
        "系统当前重点解决三个问题：一是海量日志与 Trap 中的重点风险识别；二是通过受控工具链为 AI 提供证据，避免模型凭空判断；"
        "三是让值班人员能够把经验规则写成自然语言，并在后续分析中被解释、审计和安全应用。"
    )
    doc.add_heading("1.1 建设目标", level=2)
    add_bullets(
        doc,
        [
            "统一接入 Syslog、SNMP Trap 和聚合告警事件，形成可追溯的数据底座。",
            "对 4 小时、12 小时、24 小时等时间窗口生成 AI 分析结果。",
            "提供对外演示型 AI 分析结果页，突出系统状态、核心风险和 AI 运维能力。",
            "支持 AI 分析规则：用户输入自然语言，系统调用大模型解析成结构化规则，经后端校验后保存。",
            "默认优先使用内网大模型服务；内网异常时自动 fallback 到公网模型。",
            "保留原始规则、AI解析结果、命中次数和安全例外，支持审计与回滚。",
        ],
    )

    doc.add_heading("2. 总体架构", level=1)
    doc.add_paragraph(
        "系统采用“数据接入层、存储层、事件聚合层、AI分析层、Web展示层”的分层架构。"
        "高频日志和事件数据存入 Elasticsearch，用户、任务、AI运行记录、发现项、规则和审计信息存入 MySQL。"
        "AI 不直接访问数据库，而是通过后端受控工具获取证据。"
    )
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(IMG_PATH), width=Inches(7.2))
    caption = doc.add_paragraph("图 1  宁智网维 AIOps 总体架构图")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("2.1 分层说明", level=2)
    add_table(
        doc,
        ["层级", "组件", "职责"],
        [
            ["数据接入层", "Logstash Syslog / Trap Pipeline", "监听 UDP 10087/10086，接收设备 Syslog 和 Trap，完成基础解析并写入 ES。"],
            ["数据存储层", "Elasticsearch", "保存 syslog raw、syslog parsed、trap raw、alarm events 等高频时序数据。"],
            ["元数据层", "MySQL", "保存用户、任务、AI运行记录、发现项、反馈、AI规则、命中记录、审计日志。"],
            ["事件聚合层", "aiops-event-worker", "持续读取 parsed Syslog，聚合成 alarm_events，形成分析候选。"],
            ["AI分析层", "aiops-api / aiops-scheduler / light_agent", "构建窗口摘要、调用受控工具、请求 LLM、保存结构化结果。"],
            ["模型调用层", "aiops.llm.client", "统一封装内网模型优先、公网 fallback、超时、健康检查和脱敏日志。"],
            ["展示层", "Vue + Nginx", "提供实时事件、AI分析结果、规则管理、历史、任务和系统管理页面。"],
        ],
    )

    doc.add_heading("3. 核心数据流", level=1)
    add_numbered(
        doc,
        [
            "城域网设备将 Syslog 发送到 UDP 10087，将 SNMP Trap 发送到 UDP 10086。",
            "Logstash 接收并解析数据，原始数据写入 raw 索引，结构化 Syslog 写入 parsed 索引。",
            "event-worker 周期性读取 parsed Syslog，按设备、对象、事件类型、恢复状态等维度生成 alarm_events。",
            "AI分析任务读取当前窗口 summary，并根据候选事件调用受控调查工具补充证据。",
            "AI Agent 基于 summary 与工具结果生成结构化 JSON。",
            "后端保存 AI运行记录、发现项和报告文件，前端展示为 AI分析结果页。",
            "用户创建 AI分析规则后，后端调用 LLM 解析规则并做安全校验，后续分析前将规则用于评分、降噪和报告说明。",
        ],
    )

    doc.add_heading("4. AI分析流程设计", level=1)
    doc.add_paragraph(
        "AI分析采用轻量 Agent 流程。系统先由后端构建 compact summary，再将 summary 与受控工具 schema 交给模型。"
        "模型只能请求工具名和参数，不能直接访问 Elasticsearch 或 MySQL，也不能执行自由 SQL/DSL。"
    )
    doc.add_heading("4.1 以 12 小时窗口为例", level=2)
    add_table(
        doc,
        ["步骤", "系统动作", "中间产物/输出", "交互对象"],
        [
            ["1", "确定分析窗口，例如最近 12 小时。", "window_start、window_end、hours=12。", "前端手动触发或 scheduler 定时触发。"],
            ["2", "读取 ES 中的 syslog、trap、alarm_events，生成窗口摘要。", "current_window_summary JSON。", "后端 summary builder。"],
            ["3", "加载 enabled=true 的 AI分析规则。", "规则列表、结构化 parsed_rule、安全例外配置。", "MySQL ai_analysis_rules。"],
            ["4", "对候选事件做规则预处理。", "score 调整、rule_hit 标记、suppressed 标记、安全例外标记。", "规则引擎。"],
            ["5", "第一次 LLM 调用。", "模型判断是否需要工具；常见返回 tool_calls。", "统一 LLM 客户端，优先内网模型。"],
            ["6", "执行受控工具取证。", "investigate_candidates / get_related_events / get_device_history / get_topology_context 等工具结果。", "后端工具层访问 ES。"],
            ["7", "第二次或后续 LLM 调用。", "结构化 AI分析 JSON。", "LLM + 后端 JSON 校验。"],
            ["8", "持久化结果。", "ai_analysis_runs、ai_findings、报告 JSON/Markdown、trajectory。", "MySQL + /data/jscn-aiops/reports。"],
            ["9", "前端展示。", "AI分析结果页：系统状态、AI结论、Tabs、右侧摘要、导出 Markdown。", "Vue Web 控制台。"],
        ],
    )

    doc.add_heading("4.2 真实示例说明", level=2)
    add_table(
        doc,
        ["分类", "示例输出", "说明"],
        [
            ["必须处理", "CB-CR16K-M-A XGE12/0/3 光模块 OPTICAL_FAULT ErrorCode=600005 未恢复。", "持续 open，历史有自愈但本次超时未恢复，建议排查光功率和模块状态。"],
            ["重点关注/关联分析", "CN16K-F-HeXinA 与 YGM-16K-M-A 多链路丢包告警与 BFD 联动。", "同窗口多链路质量波动，存在跨链路关联风险。"],
            ["建议动作", "排查光模块状态并准备更换；检查链路丢包根因；修正规则降噪逻辑。", "建议动作可复制，也可作为工单描述。"],
            ["规则与降噪", "RADIUS 忽略类规则若触发安全例外，仍输出说明。", "避免用户规则无条件屏蔽重大公共服务故障。"],
            ["证据不足", "Trap OID 未翻译、MIB 缺失、拓扑上下文不足。", "明确告诉值班人员哪些数据还需要补充。"],
        ],
    )

    doc.add_heading("5. AI分析规则功能设计", level=1)
    doc.add_paragraph(
        "AI分析规则用于把值班经验变成可审计、可解释、可安全应用的规则。用户在页面输入自然语言后，"
        "系统先调用 LLM 解析为结构化 JSON，然后后端执行 validateAndNormalizeRule 做兜底修正和安全校验，用户确认后才保存启用。"
    )
    add_table(
        doc,
        ["字段", "说明"],
        [
            ["raw_text", "用户原始输入，必须保留，便于审计和回滚。"],
            ["rule_type", "attention、noise_reduction、threshold、report_preference、unknown。"],
            ["action", "boost_priority、downgrade_or_suppress、threshold_control、format_control、unknown。"],
            ["target_event_families", "事件族，例如 OPTICAL_FAULT、RADIUS_ACCOUNTING_FAILURE、BFD_SESSION_FLAP。"],
            ["target_keywords / devices / objects", "关键词、设备、对象范围。"],
            ["parsed_rule", "LLM解析结果和后端归一化结果。"],
            ["priority", "规则优先级；降噪类 RADIUS 规则兜底为 30。"],
            ["safety_exceptions", "降噪规则必须包含安全例外，避免屏蔽重大故障。"],
            ["hit_count / last_hit_at", "命中次数和最近命中时间。"],
        ],
    )
    doc.add_heading("5.1 规则解析安全策略", level=2)
    add_bullets(
        doc,
        [
            "“不管、不用管、忽略、忽视、不关注、不报、不提示、降噪、屏蔽、过滤、压低”等表达强制归为降噪类，除非出现反向表达。",
            "降噪类规则必须自动补齐安全例外：多设备同源、持续未恢复超过 2 小时、严重等级、影响认证或计费核心服务。",
            "RADIUS、光模块、BFD、丢包、PPP、Trap/MIB 等关键词会自动补齐目标事件族。",
            "低置信度或 unknown 类型不允许直接启用，前端提示用户修改描述或人工确认。",
            "规则命中后，报告中需要解释哪些结论受用户规则影响。",
        ],
    )

    doc.add_heading("6. 模型调用设计", level=1)
    doc.add_paragraph(
        "系统新增统一 LLM 客户端，所有 AI分析和规则解析都走同一入口。调用策略是：内网模型启用时优先调用内网 OpenAI-compatible API；"
        "内网不可用、超时、HTTP异常或返回异常时，自动 fallback 到公网模型。"
    )
    add_table(
        doc,
        ["配置项", "说明"],
        [
            ["INTERNAL_LLM_ENABLED", "是否启用内网模型。"],
            ["INTERNAL_LLM_BASE_URL", "内网模型 OpenAI-compatible Base URL。"],
            ["INTERNAL_LLM_MODEL", "内网模型名称，当前为 deepseek-v4-pro。"],
            ["INTERNAL_LLM_API_KEY / INTERNAL_LLM_API_KEYS", "内网模型 Key，支持多个 Key 逗号分隔随机选择。"],
            ["INTERNAL_LLM_TIMEOUT", "内网请求超时，默认 30 秒。"],
            ["PUBLIC_LLM_ENABLED / PUBLIC_LLM_BASE_URL / PUBLIC_LLM_MODEL / PUBLIC_LLM_API_KEY", "公网 fallback 配置。"],
            ["AI_RULE_LLM_ENABLED", "规则解析是否启用 LLM。"],
        ],
    )
    doc.add_paragraph("安全要求：API Key 不写入代码，不提交到 GitHub，不在日志中打印；日志只记录 provider、model、耗时和错误摘要。")

    doc.add_heading("7. Web 页面使用说明", level=1)
    doc.add_paragraph("访问地址：")
    add_code(doc, "http://<服务器IP>:5772/")
    doc.add_paragraph(
        "当前页面定位为“江苏有线南京分公司城域网智能运维分析演示系统”。左侧菜单保留实时 Events、实时 Syslog、Trap 管理、"
        "AI分析结果、AI分析历史、AI分析规则、定时任务、简洁概览和系统管理等入口。"
    )
    doc.add_heading("7.1 登录", level=2)
    add_numbered(doc, ["打开 Web 地址。", "输入管理员或查看用户账号密码。", "登录后进入 AI分析结果页或默认工作台。", "登录页面不提供公开注册；用户由系统内用户管理功能维护。"])
    doc.add_heading("7.2 AI分析结果页", level=2)
    add_bullets(
        doc,
        [
            "Hero 区只展示系统核心状态、AI总结标题、一句话 AI 提示和机器人助手。",
            "Tabs 显示数量统计：必须处理、重点关注、关联分析、已恢复、建议动作、规则与降噪、证据不足。",
            "右侧栏保留 AI分析信息、恢复/噪声摘要、规则与降噪摘要。",
            "支持重新分析、查看历史和导出 Markdown。",
            "详细证据、缺失数据和建议动作放在 Tabs 内，避免首屏堆叠。",
        ],
    )
    doc.add_heading("7.3 AI分析规则", level=2)
    add_numbered(
        doc,
        [
            "进入“AI分析规则”。",
            "点击“新增规则”。",
            "输入自然语言规则，例如“光模块故障必须关注”“RADIUS所有信息都不管，忽视”“PPP认证失败超过100次才关注”。",
            "点击“解析规则”，系统调用 AI 生成结构化理解。",
            "检查系统理解、规则类型、动作、事件族、关键词、安全例外和置信度。",
            "确认无误后点击“保存并启用”。低置信度或 unknown 规则需要先修改描述。",
        ],
    )
    doc.add_heading("7.4 定时任务", level=2)
    add_bullets(doc, ["管理员可配置分析窗口、执行计划、启停状态和立即运行。", "scheduler 周期轮询到期任务，调用 AI分析流程并保存结果。", "任务执行结果可在 AI分析历史中查看。"])

    doc.add_heading("8. 部署与运维说明", level=1)
    doc.add_paragraph("部署目录：")
    add_code(doc, "/opt/jscn-aiops/deploy")
    doc.add_paragraph("数据目录：")
    add_code(doc, "/data/jscn-aiops")
    doc.add_paragraph("启动命令：")
    add_code(doc, "cd /opt/jscn-aiops/deploy\ndocker-compose up -d")
    doc.add_paragraph("常用检查命令：")
    add_code(
        doc,
        "docker-compose ps\n"
        "docker-compose logs --tail=100 aiops-api\n"
        "docker-compose logs --tail=100 aiops-event-worker\n"
        "docker-compose logs --tail=100 aiops-scheduler\n"
        "curl -s http://127.0.0.1:5772/api/health",
    )
    add_table(
        doc,
        ["服务", "默认端口", "用途"],
        [
            ["aiops-web", "5772", "Web 控制台和 /api 反向代理。"],
            ["aiops-api", "8080", "Flask API。"],
            ["Elasticsearch", "9200", "日志、Trap 和告警事件检索。"],
            ["Kibana", "5601", "ES 查询和验证。"],
            ["MySQL", "13306", "应用元数据。"],
            ["Syslog UDP", "10087", "设备 Syslog 接入。"],
            ["SNMP Trap UDP", "10086", "设备 Trap 接入。"],
        ],
    )

    doc.add_heading("9. 主要数据表与索引", level=1)
    add_table(
        doc,
        ["类型", "名称", "说明"],
        [
            ["ES索引", "jscn-aiops-syslog-raw-*", "Syslog 原始留存。"],
            ["ES索引", "jscn-aiops-syslog-parsed-*", "结构化 Syslog。"],
            ["ES索引", "jscn-aiops-trap-raw-*", "Trap 原始及基础字段。"],
            ["ES索引", "jscn-aiops-alarm-events-*", "事件聚合结果。"],
            ["MySQL表", "ai_analysis_runs", "AI分析运行记录。"],
            ["MySQL表", "ai_findings", "AI发现项。"],
            ["MySQL表", "ai_analysis_rules", "AI分析规则。"],
            ["MySQL表", "ai_analysis_rule_hits", "规则命中记录。"],
            ["MySQL表", "report_tasks", "定时分析任务。"],
            ["MySQL表", "users / audit_logs", "用户与审计。"],
        ],
    )

    doc.add_heading("10. 安全与审计设计", level=1)
    add_bullets(
        doc,
        [
            "账号由系统内管理员维护，登录页面不提供公开注册。",
            "管理员角色可以触发分析、管理规则、管理任务和用户；普通用户以查看为主。",
            "所有 API Key 均通过环境变量或 deploy/.env 注入，不写死在代码中。",
            "AI 不直接操作数据库，不执行模型生成的 SQL/DSL。",
            "AI分析规则保留 raw_text、parsed_rule、hit_count 和 last_hit_at，便于审计。",
            "降噪规则必须保留安全例外，避免重大故障被用户规则完全屏蔽。",
        ],
    )

    doc.add_heading("11. 故障处理与回滚", level=1)
    add_table(
        doc,
        ["场景", "处理方式"],
        [
            ["Web 页面无法打开", "检查 aiops-web 容器、5772 端口、Nginx 日志和 frontend/dist 是否存在。"],
            ["API 健康检查失败", "检查 aiops-api 容器日志、MySQL/ES 连接、deploy/.env 配置。"],
            ["无实时数据", "检查 Logstash、UDP 10086/10087、防火墙、ES 索引写入情况。"],
            ["AI分析失败", "检查 aiops-api 或 aiops-scheduler 日志、LLM 配置、内网模型健康检查。"],
            ["规则解析不准确", "查看 parsed_rule 和置信度；修改自然语言描述；后端兜底会修正明显反向语义。"],
            ["页面版本需回退", "使用 /data/jscn-aiops/backups 下的 tar.gz 备份恢复对应文件后 reload/recreate 服务。"],
        ],
    )

    doc.add_heading("12. 当前版本交付清单", level=1)
    add_bullets(
        doc,
        [
            "AI分析结果页：完成对外演示型视觉升级、Hero 收敛、Tabs 数量统计、机器人气泡优化。",
            "模型调用：完成内网 deepseek-v4-pro 优先、公网 fallback、健康检查与脱敏日志。",
            "AI分析规则：完成自然语言规则 AI解析、后端校验、安全例外、命中记录和前端确认流程。",
            "数据库：新增 ai_analysis_rules、ai_analysis_rule_hits 等规则相关元数据。",
            "测试：补充规则解析测试，覆盖 RADIUS 忽略、必须关注、不能忽略、光模块、伪造模块、PPP阈值等场景。",
        ],
    )

    doc.add_heading("附录 A：关键文件路径", level=1)
    add_table(
        doc,
        ["路径", "说明"],
        [
            ["frontend/src/App.vue", "前端主页面、AI分析结果页、规则页面交互。"],
            ["frontend/src/styles.css", "视觉样式、Hero、机器人、Tabs、侧栏。"],
            ["aiops/llm/client.py", "统一 LLM 调用、内网优先、公网 fallback。"],
            ["aiops/agent/light_agent.py", "轻量 AI Agent 主流程。"],
            ["aiops/rules/analysis_rules.py", "AI分析规则解析、校验、应用、命中记录。"],
            ["app/api/analysis_rules.py", "AI分析规则 API。"],
            ["app/api/llm.py", "LLM 健康检查 API。"],
            ["deploy/docker-compose.yml", "单机 Docker Compose 部署编排。"],
            ["tests/test_analysis_rules.py", "AI规则解析测试用例。"],
        ],
    )
    doc.add_heading("附录 B：版本备注", level=1)
    doc.add_paragraph("本文档对应 2026-05-20 当前 GitHub 更新版本。后续如果改造为多节点部署、结构化 JSON 报告协议或更完整的用户权限模型，应同步更新本文档。")

    doc.save(DOC_PATH)


if __name__ == "__main__":
    build_document()
    print(DOC_PATH)
    print(IMG_PATH)
